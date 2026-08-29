#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
安全审计报告生成器（灵眸·天鉴 SpiritEye）

用法：
  python generate_report.py <report.json>                    # 生成 docx + html 报告
  python generate_report.py <report.json> --upload           # 生成并上传至 OSS
  python generate_report.py <report.json> --output DIR       # 指定输出目录
  python generate_report.py <report.json> --make-template    # 生成 Word 占位符模板

数据格式：见同目录 pikachu_report.json（下划线开头键为说明，自动忽略）。

OSS 上传：
  读取环境变量 OSS_ACCESS_KEY_ID / OSS_ACCESS_KEY_SECRET / OSS_BUCKET / OSS_ENDPOINT，
  上传至 {OSS_BUCKET}.{OSS_ENDPOINT}/{OSS_KEY_PREFIX}/{date}_{project}_安全报告.{docx,html}
"""

import argparse
import html as html_mod
import json
import os
import sys
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# ---------------- 常量 ----------------
# OSS 配置从环境变量读取，禁止硬编码（防止仓库泄露）
OSS_BUCKET = os.environ.get("OSS_BUCKET", "").strip()
OSS_ENDPOINT = os.environ.get("OSS_ENDPOINT", "").strip()
OSS_KEY_PREFIX = os.environ.get("OSS_KEY_PREFIX", "reports").strip()

SEVERITY_ORDER = ["Critical", "High", "Medium", "Low", "Info"]
SEVERITY_CN = {"Critical": "严重", "High": "高危", "Medium": "中危", "Low": "低危", "Info": "提示"}
SEVERITY_COLORS = {
    "Critical": "C00000",   # 深红
    "High": "E74C3C",       # 红
    "Medium": "F39C12",     # 橙
    "Low": "2980B9",        # 蓝
    "Info": "7F8C8D",       # 灰
}
SEVERITY_COLORS_HTML = {k: "#" + v for k, v in SEVERITY_COLORS.items()}

esc = html_mod.escape


# ============================================================
# docx 辅助函数
# ============================================================
def _set_run(run, size=10.5, bold=False, color=None, mono=False, east_asia="微软雅黑"):
    """设置 run 字体：中文字体 east_asia，西文/代码 Consolas，可加粗与颜色。"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "Consolas" if mono else "Calibri"
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), east_asia)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def _add_para(doc, text, size=10.5, bold=False, color=None, mono=False,
              align=None, space_before=0, space_after=4, indent=None):
    """添加段落，返回 Paragraph。"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if align is not None:
        pf.alignment = align
    if indent is not None:
        pf.left_indent = Cm(indent)
    if text:
        run = p.add_run(text)
        _set_run(run, size=size, bold=bold, color=color, mono=mono)
    return p


def _shade_paragraph(p, fill="F2F2F2"):
    """给段落加浅灰底纹（用于代码/证据块）。"""
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def _shade_cell(cell, fill="2F3640"):
    """给表格单元格加底纹（表头深色白字）。"""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _repeat_table_header(row):
    """跨页重复表头行。"""
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def _set_cell_text(cell, text, size=9, bold=False, color=None, mono=False,
                   align=WD_ALIGN_PARAGRAPH.CENTER):
    """填充单元格文本（清空默认段落）。"""
    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    if text:
        run = p.add_run(text)
        _set_run(run, size=size, bold=bold, color=color, mono=mono)


# ============================================================
# docx 报告生成
# ============================================================
def make_docx(data, out_path):
    """按统一安全报告样式生成 Word 报告。"""
    doc = Document()

    # 默认样式与页边距
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    for section in doc.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.4)
        section.right_margin = Cm(2.4)

    title = data.get("title", "安全自查报告")
    subtitle = data.get("subtitle", "")
    date = data.get("date", "")

    # ---------- 封面 ----------
    _add_para(doc, "", space_after=24)
    _add_para(doc, title, size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
    if subtitle:
        _add_para(doc, subtitle, size=14, bold=True, color="C00000",
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    if date:
        _add_para(doc, f"审计日期：{date}", size=12,
                  align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    _add_para(doc, "", space_after=8)

    # ---------- 一 测试目标 ----------
    _add_para(doc, "一 测试目标", size=16, bold=True, space_before=12, space_after=6)
    for t in data.get("targets", []):
        name = t.get("name", "")
        desc = t.get("desc", "")
        _add_para(doc, f"  {name}（{desc}）" if desc else f"  {name}", indent=0.4)

    # ---------- 二 测试方式 ----------
    _add_para(doc, "二 测试方式", size=16, bold=True, space_before=12, space_after=6)
    _add_para(doc, "  " + "、".join(data.get("methods", [])), indent=0.4)

    # ---------- 三 漏洞统计 ----------
    _add_para(doc, "三 漏洞统计", size=16, bold=True, space_before=12, space_after=6)
    summary = data.get("summary", {})

    _add_para(doc, "3.1 总体数量", size=14, bold=True, space_after=4)
    total = summary.get("total", {})
    source_split = summary.get("source_split", {})
    rows = [
        ["严重性", "总数", "渗透测试", "代码审计"],
    ]
    for sev in SEVERITY_ORDER:
        cnt = total.get(sev, 0)
        split = source_split.get(sev, {})
        rows.append([
            f"{sev} {SEVERITY_CN.get(sev, '')}",
            str(cnt),
            str(split.get("pentest", "")),
            str(split.get("codeaudit", "")),
        ])
    rows.append([
        "总计",
        str(summary.get("total_count", sum(total.values()))),
        str(summary.get("pentest", "")),
        str(summary.get("codeaudit", "")),
    ])
    table = doc.add_table(rows=len(rows), cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            if i == 0:
                _set_cell_text(cell, val, size=10, bold=True, color="FFFFFF")
                _shade_cell(cell, "2F3640")
            elif i == len(rows) - 1:
                _set_cell_text(cell, val, size=10, bold=True)
            else:
                _set_cell_text(cell, val, size=10)

    p0 = summary.get("p0", [])
    p1 = summary.get("p1", [])
    if p0:
        _add_para(doc, "3.2 P0 级（必须修复+优先整改）", size=14, bold=True,
                  space_before=10, space_after=4)
        for item in p0:
            _add_para(doc, f"  • {item}", indent=0.4, space_after=2)
    if p1:
        _add_para(doc, "3.3 P1 级（必须修复）", size=14, bold=True,
                  space_before=10, space_after=4)
        for item in p1:
            _add_para(doc, f"  • {item}", indent=0.4, space_after=2)

    # ---------- 四 隐患详情 ----------
    doc.add_page_break()
    _add_para(doc, "四 隐患详情", size=16, bold=True, space_after=6)
    _add_finding_items(doc, data)

    # ---------- 附录：整改跟踪表 ----------
    tracking = data.get("tracking", [])
    if tracking:
        doc.add_page_break()
        _add_para(doc, "附录：整改跟踪表", size=16, bold=True, space_after=6)
        headers = ["漏洞级别", "序号", "漏洞名称", "漏洞详情", "问题编号", "修复详情", "完成情况", "复测通过情况"]
        t_table = doc.add_table(rows=1 + len(tracking), cols=8)
        t_table.style = "Table Grid"
        t_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for j, h in enumerate(headers):
            cell = t_table.cell(0, j)
            _set_cell_text(cell, h, size=9, bold=True, color="FFFFFF")
            _shade_cell(cell, "2F3640")
        _repeat_table_header(t_table.rows[0])
        for i, tr in enumerate(tracking, start=1):
            vals = [tr.get(k, "") for k in ("level", "seq", "name", "detail", "id", "fix", "status", "retest")]
            for j, val in enumerate(vals):
                align = WD_ALIGN_PARAGRAPH.LEFT if j in (2, 3, 5) else WD_ALIGN_PARAGRAPH.CENTER
                _set_cell_text(t_table.cell(i, j), str(val), size=9, align=align)

    doc.save(out_path)


def _add_finding_items(doc, data):
    """按章节与编号规则输出全部漏洞条目。"""
    findings = data.get("findings", [])
    sections = data.get("sections", [])
    for sec in sections:
        sec_id = sec["id"]
        sec_findings = [f for f in findings if str(f.get("section", sec_id)) == str(sec_id)]
        if not sec_findings:
            continue
        _add_para(doc, f"{sec_id} {sec['title']}", size=14, bold=True,
                  space_before=12, space_after=4)
        if sec.get("note"):
            _add_para(doc, sec["note"], size=10, color="595959", space_after=4)

        if sec.get("group_by_severity"):
            sub = 1
            for sev in SEVERITY_ORDER:
                group = [f for f in sec_findings if f.get("severity") == sev]
                if not group:
                    continue
                _add_para(doc, f"{sec_id}.{sub} {sev}（{len(group)} 个）", size=12,
                          bold=True, space_before=8, space_after=4)
                for f in group:
                    _add_finding(doc, None, f)
                sub += 1
        else:
            for i, f in enumerate(sec_findings, 1):
                _add_finding(doc, f"{sec_id}.{i}", f)


def _add_finding(doc, num, f):
    """渲染单个漏洞条目。"""
    head = f"{num} {f.get('id', '')} {f.get('title', '')}".strip()
    _add_para(doc, head, size=13, bold=True, space_before=10, space_after=2)

    # 严重性行（severity 着色）
    sev = f.get("severity", "")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("严重性：")
    _set_run(run, bold=True)
    run = p.add_run(sev)
    _set_run(run, bold=True, color=SEVERITY_COLORS.get(sev, "000000"))
    meta = f" | OWASP：{f.get('owasp', '')} | STRIDE：{f.get('stride', '')}"
    if f.get("status"):
        meta += f" | 状态：{f['status']}"
    run = p.add_run(meta)
    _set_run(run)

    # 漏洞描述
    desc = f.get("description", "")
    if desc:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("漏洞描述：")
        _set_run(run, bold=True)
        run = p.add_run(desc)
        _set_run(run)

    # 证据块
    for ev in f.get("evidence", []):
        label = ev.get("label", "")
        content = str(ev.get("content", "")).split("\n")
        if label:
            _add_para(doc, label, size=10.5, bold=True, space_before=4, space_after=2)
        if ev.get("code"):
            for line in content:
                if line.strip() or line != content[-1]:
                    cp = _add_para(doc, line if line else " ", size=9, mono=True,
                                   space_after=0, indent=0.3)
                    _shade_paragraph(cp)
        else:
            for line in content:
                _add_para(doc, line if line else " ", size=10.5, space_after=1, indent=0.3)

    # 实际执行修复方案
    actual_fix = f.get("actual_fix", "")
    if actual_fix:
        _add_para(doc, "实际执行修复方案：", size=10.5, bold=True,
                  space_before=4, space_after=2)
        for line in str(actual_fix).split("\n"):
            _add_para(doc, line if line else " ", size=10.5, space_after=1, indent=0.3)

    # 参考
    ref = f.get("reference", "")
    if ref:
        p = _add_para(doc, f"参考：{ref}", size=9, color="595959", space_before=2, space_after=2)
        _shade_paragraph(p, "F5F5F5")


# ============================================================
# docx 占位符模板生成
# ============================================================
def build_template_docx(out_path):
    """生成带 {{占位符}} 的 Word 模板，供人工参照格式。"""
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    _add_para(doc, "", space_after=24)
    _add_para(doc, "{{title}}", size=22, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
    _add_para(doc, "{{subtitle}}", size=14, bold=True, color="C00000",
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    _add_para(doc, "审计日期：{{date}}", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    _add_para(doc, "", space_after=8)
    _add_para(doc, "【使用说明】本文件为报告格式模板，正式报告由 generate_report.py 从 JSON 自动生成："
                   "python generate_report.py <项目>_report.json --upload", size=9, color="595959")

    _add_para(doc, "一 测试目标", size=16, bold=True, space_before=12, space_after=6)
    _add_para(doc, "  {{target.1.name}}（{{target.1.desc}}）", indent=0.4)
    _add_para(doc, "  {{target.2.name}}（{{target.2.desc}}）", indent=0.4)

    _add_para(doc, "二 测试方式", size=16, bold=True, space_before=12, space_after=6)
    _add_para(doc, "  {{methods}}", indent=0.4)

    _add_para(doc, "三 漏洞统计", size=16, bold=True, space_before=12, space_after=6)
    _add_para(doc, "3.1 总体数量", size=14, bold=True, space_after=4)
    table = doc.add_table(rows=6, cols=4)
    table.style = "Table Grid"
    headers = ["严重性", "总数", "渗透测试", "代码审计"]
    rows = [["Critical 严重", "{{n_critical}}", "", ""],
            ["High 高危", "{{n_high}}", "", ""],
            ["Medium 中危", "{{n_medium}}", "", ""],
            ["Low 低危", "{{n_low}}", "", ""],
            ["总计", "{{n_total}}", "", ""]]
    for j, h in enumerate(headers):
        _set_cell_text(table.cell(0, j), h, size=10, bold=True, color="FFFFFF")
        _shade_cell(table.cell(0, j))
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            _set_cell_text(table.cell(i, j), val, size=10)
    _add_para(doc, "3.2 P0 级（必须修复+优先整改）", size=14, bold=True, space_before=10, space_after=4)
    _add_para(doc, "  • {{p0.1}}", indent=0.4, space_after=2)
    _add_para(doc, "  • {{p0.2}}", indent=0.4, space_after=2)
    _add_para(doc, "3.3 P1 级（必须修复）", size=14, bold=True, space_before=10, space_after=4)
    _add_para(doc, "  • {{p1.1}}", indent=0.4, space_after=2)
    _add_para(doc, "  • {{p1.2}}", indent=0.4, space_after=2)

    doc.add_page_break()
    _add_para(doc, "四 隐患详情", size=16, bold=True, space_after=6)
    _add_para(doc, "1 {{section.1.title}}", size=14, bold=True, space_before=12, space_after=4)
    _add_para(doc, "1.1 {{finding.id}} {{finding.title}}", size=13, bold=True, space_before=10, space_after=2)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("严重性：")
    _set_run(run, bold=True)
    run = p.add_run("{{finding.severity}}")
    _set_run(run, bold=True, color="C00000")
    run = p.add_run(" | OWASP：{{finding.owasp}} | STRIDE：{{finding.stride}}")
    _set_run(run)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("漏洞描述：")
    _set_run(run, bold=True)
    run = p.add_run("{{finding.description}}")
    _set_run(run)
    _add_para(doc, "渗透证据（{{evidence.id}}）：", size=10.5, bold=True, space_before=4, space_after=2)
    for line in ["{{evidence.request}}", "{{evidence.response}}"]:
        cp = _add_para(doc, line, size=9, mono=True, space_after=0, indent=0.3)
        _shade_paragraph(cp)
    _add_para(doc, "利用方式：", size=10.5, bold=True, space_before=4, space_after=2)
    _add_para(doc, "  {{finding.exploitation}}", indent=0.3)
    _add_para(doc, "修复建议：", size=10.5, bold=True, space_before=4, space_after=2)
    _add_para(doc, "  {{finding.remediation}}", indent=0.3)
    _add_para(doc, "2 {{section.2.title}}", size=14, bold=True, space_before=12, space_after=4)
    _add_para(doc, "2.1 {{severity.group}}（{{n}} 个）", size=12, bold=True, space_before=8, space_after=4)
    _add_para(doc, "{{finding.id}} {{finding.title}}", size=13, bold=True, space_before=10, space_after=2)

    doc.add_page_break()
    _add_para(doc, "附录：整改跟踪表", size=16, bold=True, space_after=6)
    headers = ["漏洞级别", "序号", "漏洞名称", "漏洞详情", "问题编号", "修复详情", "完成情况", "复测通过情况"]
    t_table = doc.add_table(rows=4, cols=8)
    t_table.style = "Table Grid"
    for j, h in enumerate(headers):
        _set_cell_text(t_table.cell(0, j), h, size=9, bold=True, color="FFFFFF")
        _shade_cell(t_table.cell(0, j))
    _repeat_table_header(t_table.rows[0])
    for i in range(1, 4):
        for j in range(8):
            _set_cell_text(t_table.cell(i, j), "", size=9)

    doc.save(out_path)


# ============================================================
# HTML 报告生成
# ============================================================
HTML_TEMPLATE = """<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>__TITLE__</title>
<style>
:root{--critical:#c0392b;--high:#e74c3c;--medium:#f39c12;--low:#2980b9;--info:#7f8c8d}
*{box-sizing:border-box}
body{font-family:"Microsoft YaHei","PingFang SC","Helvetica Neue",Arial,sans-serif;background:#f0f2f5;color:#2c3e50;margin:0;line-height:1.6}
.container{max-width:1100px;margin:0 auto;padding:20px 16px 60px}
.cover{background:linear-gradient(135deg,#1b2a4a 0%,#0f2027 100%);color:#fff;padding:44px 40px;border-radius:12px;margin-bottom:20px}
.cover h1{margin:0 0 10px;font-size:28px;letter-spacing:1px}
.cover .subtitle{color:#ff6b6b;font-size:16px;font-weight:600;margin-bottom:8px}
.cover .meta{opacity:.85;font-size:13px}
.card{background:#fff;border-radius:12px;padding:20px 24px;margin-bottom:16px;box-shadow:0 1px 3px rgba(0,0,0,.08)}
.card h2{margin:0 0 10px;font-size:20px;border-left:4px solid #1b2a4a;padding-left:10px}
.card h3{margin:18px 0 8px;font-size:16px}
.meta-row{font-size:13px;color:#555;margin-bottom:4px}
.stats{display:flex;gap:12px;flex-wrap:wrap;margin:12px 0}
.stat-card{flex:1;min-width:130px;background:#fff;border-radius:10px;padding:14px 10px;text-align:center;box-shadow:0 1px 3px rgba(0,0,0,.08);border-top:4px solid var(--c)}
.stat-card .num{font-size:30px;font-weight:700;color:var(--c)}
.stat-card .label{font-size:13px;color:#666;margin-top:2px}
.toolbar{position:sticky;top:0;z-index:10;background:#fff;border-radius:10px;padding:12px 16px;margin-bottom:16px;box-shadow:0 2px 8px rgba(0,0,0,.1);display:flex;gap:10px;align-items:center;flex-wrap:wrap}
.toolbar input,.toolbar select{padding:7px 12px;border:1px solid #dcdde1;border-radius:6px;font-size:14px;font-family:inherit}
.toolbar input{flex:1;min-width:220px}
.badge{display:inline-block;padding:2px 12px;border-radius:12px;color:#fff;font-size:12px;font-weight:700;margin-right:8px;vertical-align:2px}
.badge.Critical{background:var(--critical)}.badge.High{background:var(--high)}.badge.Medium{background:var(--medium)}.badge.Low{background:var(--low)}.badge.Info{background:var(--info)}
.finding{background:#fff;border-radius:10px;padding:18px 22px;margin-bottom:14px;box-shadow:0 1px 3px rgba(0,0,0,.08);border-left:4px solid #dcdde1}
.finding.sev-Critical{border-left-color:var(--critical)}.finding.sev-High{border-left-color:var(--high)}
.finding.sev-Medium{border-left-color:var(--medium)}.finding.sev-Low{border-left-color:var(--low)}.finding.sev-Info{border-left-color:var(--info)}
.f-head{display:flex;align-items:baseline;gap:10px;flex-wrap:wrap}
.f-head h3{margin:0;font-size:17px;flex:1;min-width:200px}
.f-id{font-family:Consolas,monospace;font-size:13px;color:#888}
.f-meta{font-size:13px;color:#555;margin:6px 0}
.f-meta b{color:#2c3e50}
.f-desc{margin:8px 0;font-size:14px}
.ev-label{font-weight:600;font-size:14px;margin:12px 0 4px}
pre{background:#f6f8fa;border:1px solid #e1e4e8;border-radius:6px;padding:12px;overflow-x:auto;font-family:Consolas,"Courier New",monospace;font-size:12.5px;line-height:1.5;margin:4px 0;white-space:pre-wrap;word-break:break-all}
.p-text{font-size:14px;margin:2px 0}
.actual-fix{background:#fffbe6;border:1px solid #ffe58f;border-radius:6px;padding:10px 14px;font-size:13.5px;margin-top:10px;white-space:pre-wrap}
table{width:100%;border-collapse:collapse;background:#fff;margin-top:8px}
th,td{border:1px solid #dcdde1;padding:7px 8px;font-size:13px;text-align:left;vertical-align:top}
th{background:#2f3640;color:#fff;font-weight:600;white-space:nowrap}
tr:nth-child(even) td{background:#fafbfc}
details{border:1px solid #e1e4e8;border-radius:6px;margin:8px 0;background:#fafbfc}
details summary{padding:8px 12px;cursor:pointer;font-weight:600;font-size:14px}
details .ev-inner{padding:0 12px 12px}
.f-details{border:none;border-top:1px dashed #dcdde1;border-radius:0;margin:10px 0 0;background:transparent}
.f-details summary{padding:8px 0;font-size:14px;color:#555;user-select:none}
.f-details summary:hover{color:#1b2a4a}
.btn{padding:7px 12px;border:1px solid #dcdde1;border-radius:6px;background:#fff;font-size:13px;font-family:inherit;cursor:pointer;color:#555}
.btn:hover{border-color:#1b2a4a;color:#1b2a4a}
footer{text-align:center;color:#999;font-size:12px;margin-top:30px}
@media print{
  body{background:#fff}
  .toolbar{display:none}
  .cover{background:#fff;color:#000;border:2px solid #1b2a4a}
  .cover .subtitle{color:#c00000}
  .finding,.card{box-shadow:none;border:1px solid #ccc;page-break-inside:avoid}
  .container{max-width:100%;padding:0}
}
</style>
</head>
<body>
<div class="container">
  <header class="cover">
    <h1>__TITLE__</h1>
    <div class="subtitle">__SUBTITLE__</div>
    <div class="meta">审计日期：__DATE__　|　测试方式：__METHODS__</div>
  </header>

  <div class="toolbar">
    <input id="search" type="text" placeholder="搜索漏洞编号 / 名称 / 描述...">
    <select id="sevFilter">
      <option value="">全部严重性</option>
      <option value="Critical">Critical 严重</option>
      <option value="High">High 高危</option>
      <option value="Medium">Medium 中危</option>
      <option value="Low">Low 低危</option>
      <option value="Info">Info 提示</option>
    </select>
    <button id="expandAll" class="btn" type="button">全部展开</button>
    <button id="collapseAll" class="btn" type="button">全部折叠</button>
    <span id="count" style="font-size:13px;color:#888"></span>
  </div>

  <section class="card" id="targets">
    <h2>一 测试目标</h2>
    __TARGETS__
  </section>

  <section class="card" id="summary">
    <h2>三 漏洞统计</h2>
    <div class="stats">__STAT_CARDS__</div>
    <h3>3.1 总体数量</h3>
    <table>
      <tr><th>严重性</th><th>总数</th><th>渗透测试</th><th>代码审计</th></tr>
      __SUMMARY_ROWS__
    </table>
    <h3>3.2 P0 级（必须修复+优先整改）</h3>
    <ul>__P0_LIST__</ul>
    <h3>3.3 P1 级（必须修复）</h3>
    <ul>__P1_LIST__</ul>
  </section>

  <section class="card" id="findings">
    <h2>四 隐患详情</h2>
    __FINDINGS__
  </section>

  <section class="card" id="tracking">
    <h2>附录：整改跟踪表</h2>
    <table>
      <tr><th>漏洞级别</th><th>序号</th><th>漏洞名称</th><th>漏洞详情</th><th>问题编号</th><th>修复详情</th><th>完成情况</th><th>复测通过情况</th></tr>
      __TRACKING_ROWS__
    </table>
  </section>

  <footer>由灵眸·天鉴（SpiritEye）安全审计工具生成</footer>
</div>
<script>
(function(){
  var search = document.getElementById('search');
  var sevSel = document.getElementById('sevFilter');
  var countEl = document.getElementById('count');
  var expandAll = document.getElementById('expandAll');
  var collapseAll = document.getElementById('collapseAll');
  function apply(){
    var q = (search.value || '').toLowerCase().trim();
    var sev = sevSel.value;
    var cards = document.querySelectorAll('.finding');
    var shown = 0;
    cards.forEach(function(card){
      var ok = true;
      if(sev && card.getAttribute('data-severity') !== sev) ok = false;
      if(ok && q && card.textContent.toLowerCase().indexOf(q) < 0) ok = false;
      card.style.display = ok ? '' : 'none';
      if(ok) shown++;
    });
    countEl.textContent = '显示 ' + shown + ' / ' + cards.length + ' 个漏洞';
  }
  search.addEventListener('input', apply);
  sevSel.addEventListener('change', apply);
  expandAll.addEventListener('click', function(){
    document.querySelectorAll('.f-details').forEach(function(d){ d.open = true; });
  });
  collapseAll.addEventListener('click', function(){
    document.querySelectorAll('.f-details').forEach(function(d){ d.open = false; });
  });
  apply();
})();
</script>
</body>
</html>
"""


def _finding_html(num, f):
    """渲染单个漏洞 HTML 卡片。"""
    sev = f.get("severity", "")
    sev_cn = SEVERITY_CN.get(sev, "")
    parts = [
        f'<div class="finding sev-{esc(sev)}" id="{esc(f.get("id", ""))}" data-severity="{esc(sev)}">',
        f'<div class="f-head"><span class="badge {esc(sev)}">{esc(sev)} {esc(sev_cn)}</span>'
        f'<span class="f-id">{esc(f.get("id", ""))}</span>'
        f'<h3>{esc(f.get("title", ""))}</h3></div>',
        f'<div class="f-meta"><b>OWASP：</b>{esc(f.get("owasp", ""))}　<b>STRIDE：</b>{esc(f.get("stride", ""))}',
    ]
    if f.get("status"):
        parts.append(f'　<b>状态：</b>{esc(f["status"])}')
    parts.append("</div>")

    # 详情区（支持折叠）
    parts.append('<details class="f-details" open><summary>漏洞详情（点击折叠/展开）</summary>')

    if f.get("description"):
        parts.append(f'<div class="f-desc">{esc(f["description"])}</div>')

    for ev in f.get("evidence", []):
        label = ev.get("label", "")
        content = ev.get("content", "")
        if label:
            parts.append(f'<div class="ev-label">{esc(label)}</div>')
        if ev.get("code"):
            parts.append(f"<pre>{esc(str(content))}</pre>")
        else:
            for line in str(content).split("\n"):
                parts.append(f'<div class="p-text">{esc(line) if line else "&nbsp;"}</div>')

    if f.get("actual_fix"):
        parts.append(f'<div class="actual-fix"><b>实际执行修复方案：</b>\n{esc(str(f["actual_fix"]))}</div>')
    if f.get("reference"):
        parts.append(f'<div class="f-meta" style="color:#999">参考：{esc(f["reference"])}</div>')
    parts.append("</details>")
    parts.append("</div>")
    return "\n".join(parts)


def _findings_html(data):
    """按章节与编号规则输出全部漏洞 HTML。"""
    findings = data.get("findings", [])
    sections = data.get("sections", [])
    out = []
    for sec in sections:
        sec_id = sec["id"]
        sec_findings = [f for f in findings if str(f.get("section", sec_id)) == str(sec_id)]
        if not sec_findings:
            continue
        out.append(f'<h3>{esc(sec_id)} {esc(sec["title"])}</h3>')
        if sec.get("note"):
            out.append(f'<div class="f-meta">{esc(sec["note"])}</div>')
        if sec.get("group_by_severity"):
            sub = 1
            for sev in SEVERITY_ORDER:
                group = [f for f in sec_findings if f.get("severity") == sev]
                if not group:
                    continue
                out.append(f'<h3 style="font-size:15px">{esc(sec_id)}.{sub} {esc(sev)}（{len(group)} 个）</h3>')
                for f in group:
                    out.append(_finding_html(None, f))
                sub += 1
        else:
            for i, f in enumerate(sec_findings, 1):
                out.append(_finding_html(f"{sec_id}.{i}", f))
    return "\n".join(out)


def make_html(data, out_path):
    """生成自包含 HTML 报告。"""
    summary = data.get("summary", {})
    total = summary.get("total", {})

    # 统计卡片
    cards = []
    for sev in SEVERITY_ORDER:
        n = total.get(sev, 0)
        cards.append(
            f'<div class="stat-card" style="--c:{SEVERITY_COLORS_HTML[sev]}">'
            f'<div class="num">{n}</div><div class="label">{esc(sev)} {esc(SEVERITY_CN[sev])}</div></div>'
        )
    cards.append(
        f'<div class="stat-card" style="--c:#1b2a4a"><div class="num">{esc(str(summary.get("pentest", "")))}</div>'
        f'<div class="label">渗透测试</div></div>'
    )
    cards.append(
        f'<div class="stat-card" style="--c:#34495e"><div class="num">{esc(str(summary.get("codeaudit", "")))}</div>'
        f'<div class="label">代码审计</div></div>'
    )

    # 总体数量表
    source_split = summary.get("source_split", {})
    rows = []
    for sev in SEVERITY_ORDER:
        split = source_split.get(sev, {})
        rows.append(
            f'<tr><td><span class="badge {esc(sev)}">{esc(sev)}</span></td>'
            f'<td>{total.get(sev, 0)}</td><td>{split.get("pentest", "")}</td>'
            f'<td>{split.get("codeaudit", "")}</td></tr>'
        )
    rows.append(
        f'<tr><td><b>总计</b></td><td><b>{summary.get("total_count", sum(total.values()))}</b></td>'
        f'<td><b>{summary.get("pentest", "")}</b></td><td><b>{summary.get("codeaudit", "")}</b></td></tr>'
    )

    # P0 / P1 清单
    def _ul(items):
        return "".join(f"<li>{esc(str(item))}</li>" for item in items)

    targets = "".join(
        f'<div class="meta-row">• {esc(t.get("name", ""))}（{esc(t.get("desc", ""))}）</div>'
        for t in data.get("targets", [])
    )

    html = (HTML_TEMPLATE
            .replace("__TITLE__", esc(data.get("title", "安全自查报告")))
            .replace("__SUBTITLE__", esc(data.get("subtitle", "")))
            .replace("__DATE__", esc(data.get("date", "")))
            .replace("__METHODS__", esc("、".join(data.get("methods", []))))
            .replace("__TARGETS__", targets)
            .replace("__STAT_CARDS__", "\n".join(cards))
            .replace("__SUMMARY_ROWS__", "\n".join(rows))
            .replace("__P0_LIST__", _ul(summary.get("p0", [])))
            .replace("__P1_LIST__", _ul(summary.get("p1", [])))
            .replace("__FINDINGS__", _findings_html(data))
            .replace("__TRACKING_ROWS__", _tracking_html(data.get("tracking", []))))

    with open(out_path, "w", encoding="utf-8") as f:
        f.write(html)


def _tracking_html(tracking):
    rows = []
    for tr in tracking:
        rows.append(
            "<tr>"
            + "".join(
                f"<td>{esc(str(tr.get(k, '')))}</td>"
                for k in ("level", "seq", "name", "detail", "id", "fix", "status", "retest")
            )
            + "</tr>"
        )
    return "\n".join(rows)


# ============================================================
# OSS 上传
# ============================================================
def upload_to_oss(paths):
    """上传文件到 OSS bucket，key 为 {OSS_KEY_PREFIX}/<文件名>。"""
    ak_id = os.environ.get("OSS_ACCESS_KEY_ID", "").strip()
    ak_secret = os.environ.get("OSS_ACCESS_KEY_SECRET", "").strip()
    if not ak_id or not ak_secret:
        print("[ERROR] 未检测到 OSS 凭证，请先设置环境变量：")
        print("  set OSS_ACCESS_KEY_ID=<您的AccessKeyId>")
        print("  set OSS_ACCESS_KEY_SECRET=<您的AccessKeySecret>")
        sys.exit(1)
    if not OSS_BUCKET or not OSS_ENDPOINT:
        print("[ERROR] 未检测到 OSS 目标配置，请设置环境变量：")
        print("  set OSS_BUCKET=<您的bucket名称>")
        print("  set OSS_ENDPOINT=<您的endpoint，如 https://oss-cn-xxx.aliyuncs.com>")
        sys.exit(1)
    try:
        import oss2
    except ImportError:
        print("[ERROR] 缺少 oss2 依赖，请执行：pip install oss2")
        sys.exit(1)

    auth = oss2.Auth(ak_id, ak_secret)
    bucket = oss2.Bucket(auth, OSS_ENDPOINT, OSS_BUCKET)
    print(f"[OSS] bucket={OSS_BUCKET} endpoint={OSS_ENDPOINT}")
    for path in paths:
        key = f"{OSS_KEY_PREFIX}/{os.path.basename(path)}"
        bucket.put_object_from_file(key, path)
        url = f"https://{OSS_BUCKET}.{OSS_ENDPOINT.removeprefix('https://')}/{key}"
        print(f"[OSS] 已上传：{url}")


# ============================================================
# 入口
# ============================================================
def load_report(json_path):
    with open(json_path, "r", encoding="utf-8") as f:
        data = json.load(f)
    if "report" not in data:
        print("[ERROR] JSON 中缺少 report 根节点，请参照 pikachu_report.json 格式。")
        sys.exit(1)
    return data["report"]


def main():
    parser = argparse.ArgumentParser(description="安全审计报告生成器（docx + html + OSS 上传）")
    parser.add_argument("json_path", help="审计结果 JSON（参照 pikachu_report.json 格式）")
    parser.add_argument("--output", "-o", default=None, help="输出目录（默认 JSON 所在目录）")
    parser.add_argument("--upload", action="store_true", help="生成后上传至 OSS")
    parser.add_argument("--make-template", action="store_true", help="生成 Word 占位符模板并退出")
    args = parser.parse_args()

    if args.make_template:
        out = args.output or os.path.join(BASE_DIR, "report_template.docx")
        build_template_docx(out)
        print(f"[OK] Word 模板已生成：{out}")
        return

    data = load_report(args.json_path)
    date = data.get("date", datetime.now().strftime("%Y-%m-%d"))
    project = data.get("project", "项目")
    out_dir = args.output or os.path.dirname(os.path.abspath(args.json_path))
    os.makedirs(out_dir, exist_ok=True)
    base_name = f"{date}_{project}_安全自查报告"
    docx_path = os.path.join(out_dir, base_name + ".docx")
    html_path = os.path.join(out_dir, base_name + ".html")

    make_docx(data, docx_path)
    print(f"[OK] Word 报告已生成：{docx_path}")
    make_html(data, html_path)
    print(f"[OK] HTML 报告已生成：{html_path}")

    if args.upload:
        upload_to_oss([docx_path, html_path])
        print(f"[OK] 报告已上传至 OSS（{OSS_BUCKET}.{OSS_ENDPOINT}/{OSS_KEY_PREFIX}/）")


if __name__ == "__main__":
    main()
